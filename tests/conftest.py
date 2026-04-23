"""Shared test fixtures and utilities."""

from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest

from document_semantic.agents.regex_recognizer import RegexRecognizer
from document_semantic.models.semantic_document import SemanticDocument
from document_semantic.services.parsers.python_docx_parser import PythonDocxParser

FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ---------------------------------------------------------------------------
# Fixture loaders
# ---------------------------------------------------------------------------


def get_fixture_path(name: str) -> Path:
    """Resolve a fixture path relative to tests/fixtures/.

    Args:
        name: Fixture filename (e.g., 'simple_document.docx').

    Returns:
        Absolute path to the fixture file.
    """
    return FIXTURES_DIR / name


def load_expected_output(name: str) -> SemanticDocument:
    """Load expected SemanticDocument from a JSON fixture.

    Args:
        name: Fixture name without extension (e.g., 'simple_document').

    Returns:
        A SemanticDocument instance.
    """
    json_path = FIXTURES_DIR / f"{name}_expected.json"
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)
    return SemanticDocument.model_validate(data)


# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def fixtures_dir() -> Path:
    """Return the fixtures directory path."""
    return FIXTURES_DIR


@pytest.fixture
def python_docx_parser() -> PythonDocxParser:
    """Return a PythonDocxParser instance."""
    return PythonDocxParser()


@pytest.fixture
def regex_recognizer() -> RegexRecognizer:
    """Return a RegexRecognizer instance."""
    return RegexRecognizer()


@pytest.fixture
def sample_docx_path(fixtures_dir: Path, tmp_path: Path) -> Path:
    """Create a minimal DOCX file for testing using python-docx.

    Creates a simple document with title, heading, and paragraph.
    """
    from docx import Document

    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Sample Document", level=0)  # Title
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is a sample paragraph with **bold** and `code`.")
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("The end.")
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def academic_docx_path(fixtures_dir: Path, tmp_path: Path) -> Path:
    """Create an academic-style DOCX for testing.

    Contains abstract, headings, references.
    """
    from docx import Document

    doc_path = tmp_path / "academic.docx"
    doc = Document()
    doc.add_heading("Research Paper Title", level=0)
    doc.add_paragraph("Abstract")
    doc.add_paragraph(
        "This paper presents a novel approach to document semantic analysis using pipeline-based architectures."
    )
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("Background context here.")
    doc.add_heading("2. Method", level=1)
    doc.add_paragraph("Methodology description.")
    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Author A. Title of Paper. Journal, 2024.")
    doc.add_paragraph("[2] Author B. Another Paper. Conference, 2023.")
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def table_docx_path(fixtures_dir: Path, tmp_path: Path) -> Path:
    """Create a DOCX with a table for testing."""
    from docx import Document

    doc_path = tmp_path / "with_table.docx"
    doc = Document()
    doc.add_heading("Document with Table", level=0)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "10"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "20"
    doc.add_paragraph("Text after table.")
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def inline_formatting_docx_path(fixtures_dir: Path, tmp_path: Path) -> Path:
    """Create a DOCX with rich inline formatting for testing."""
    from docx import Document
    from docx.shared import Pt

    doc_path = tmp_path / "inline_formatting.docx"
    doc = Document()
    doc.add_heading("Inline Formatting Test", level=0)

    # Paragraph with bold run
    p = doc.add_paragraph()
    p.add_run("This has ")
    bold_run = p.add_run("bold text")
    bold_run.bold = True
    p.add_run(" and ")
    italic_run = p.add_run("italic text")
    italic_run.italic = True
    p.add_run(" and ")
    code_run = p.add_run("code here")
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(10)

    # Paragraph with a formula-like text
    doc.add_paragraph("The equation is E = mc^2, which is famous.")

    doc.save(str(doc_path))
    return doc_path


def _pandoc_available() -> bool:
    """Check if pandoc binary is on PATH."""
    return shutil.which("pandoc") is not None


def requires_pandoc(func):
    """Decorator to skip a test if pandoc is not installed."""
    return pytest.mark.skipif(
        not _pandoc_available(),
        reason="pandoc is not installed",
    )(func)
