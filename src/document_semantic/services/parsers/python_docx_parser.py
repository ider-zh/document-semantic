"""Python-docx parser implementation."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from document_semantic.core.logger import get_logger
from document_semantic.models.inline_elements import (
    BoldInlineElement,
    InlineElement,
    ItalicInlineElement,
)
from document_semantic.models.processor_output import ProcessorConfig, ProcessResult
from document_semantic.utils.markdown_generator import MarkdownGenerator

from .protocol import Attachment, IntermediateBlock, IntermediateResult, Parser
from .registry import ParserRegistry

logger = get_logger(__name__)

# Mapping from python-docx style names to style hints
_STYLE_HINT_MAP: dict[str, str] = {
    "Title": "Title",
    "Heading 1": "Heading1",
    "Heading 2": "Heading2",
    "Heading 3": "Heading3",
    "Heading 4": "Heading4",
    "Heading 5": "Heading5",
    "Heading 6": "Heading6",
    "List Bullet": "ListBullet",
    "List Number": "ListNumber",
    "Caption": "image_description",
    "Image Description": "image_description",
    "Table Description": "table_description",
}

# Inline element detection patterns (applied to run text)
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_CODE_RE = re.compile(r"`(.+?)`")
_LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")


def _get_style_hint(paragraph) -> str | None:
    """Extract a style hint from a python-docx paragraph.

    Also checks if the paragraph consists entirely of bold runs,
    which is a common heading pattern in documents without Heading styles.
    """
    style_name = paragraph.style.name if paragraph.style else None
    if style_name and style_name in _STYLE_HINT_MAP:
        return _STYLE_HINT_MAP[style_name]

    # Heuristic: short text where all non-empty runs are bold -> likely heading
    runs_with_text = [r for r in paragraph.runs if r.text.strip()]
    if runs_with_text and all(r.bold for r in runs_with_text):
        if len(paragraph.text.strip()) < 200:
            return "Heading1"

    return "Normal"


def _build_content_with_markdown(paragraph) -> tuple[str, list[InlineElement]]:
    """Build content string with markdown markup from runs.

    Also extracts inline elements with correct offsets.

    Returns:
        Tuple of (content_with_markdown, list_of_inline_elements)
    """
    parts: list[str] = []
    inline_elements: list[InlineElement] = []
    offset = 0

    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Detect if the run has bold or italic formatting
        prefix = ""
        suffix = ""
        if run.bold:
            prefix += "**"
            suffix = "**" + suffix
        if run.italic:
            prefix += "*"
            suffix = "*" + suffix

        wrapped_text = prefix + text + suffix
        start = offset
        end = offset + len(wrapped_text)

        # Create inline elements for this run
        if run.bold:
            b_start = start
            b_end = start + len(prefix) + len(text)
            # The bold element should cover the full wrapped region
            inline_elements.append(
                BoldInlineElement(
                    text=text,
                    start_offset=b_start,
                    end_offset=b_end,
                )
            )
        elif run.italic:
            i_start = start
            i_end = start + len(prefix) + len(text)
            inline_elements.append(
                ItalicInlineElement(
                    text=text,
                    start_offset=i_start,
                    end_offset=i_end,
                )
            )

        parts.append(wrapped_text)
        offset = end

    content = "".join(parts)
    return content, inline_elements


class PythonDocxParser(Parser):
    """Parser that uses python-docx to read DOCX structure.

    Extracts paragraphs, tables, and inline formatting information.
    """

    @property
    def name(self) -> str:
        return "python-docx"

    def parse(self, docx_path: Path, skip_image_ocr: bool = False) -> IntermediateResult:
        """Parse a DOCX file using python-docx.

        Args:
            docx_path: Absolute path to the DOCX file.

        Returns:
            IntermediateResult with blocks, metadata, and attachments.
        """
        logger.info(f"[parsing:python-docx] Opening {docx_path}")

        doc = Document(str(docx_path))
        blocks: list[IntermediateBlock] = []
        attachments: list[Attachment] = []

        # Extract document metadata
        core_properties = doc.core_properties
        metadata: dict[str, Any] = {}
        if core_properties.title:
            metadata["title"] = core_properties.title
        if core_properties.author:
            metadata["author"] = core_properties.author
        metadata["source_path"] = str(docx_path)

        # Process paragraphs
        for para in doc.paragraphs:
            content = para.text.strip()
            if not content:
                continue
            style_hint = _get_style_hint(para)

            # Build content with markdown markup from runs
            md_content, inline_elements = _build_content_with_markdown(para)
            logger.debug(
                f"[parsing:python-docx] Paragraph: style={style_hint}, "
                f"len={len(content)}, inlines={len(inline_elements)}"
            )
            blocks.append(
                IntermediateBlock(
                    content=md_content,
                    style_hint=style_hint,
                    inline_elements=inline_elements,
                )
            )

        # Process tables
        for idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cell_texts = [cell.text for cell in row.cells]
                rows.append(cell_texts)
            if rows:
                headers = rows[0] if rows else []
                table_content = f"[Table: {len(rows)} rows x {len(headers)} cols]"
                logger.debug(f"[parsing:python-docx] Table detected: {table_content}")
                blocks.append(
                    IntermediateBlock(
                        content=table_content,
                        style_hint="Table",
                    )
                )

        # Extract images from document relationships
        try:
            image_parts = doc.part.related_parts
            for rel_id, rel in image_parts.items():
                if "image" in rel.target_ref.lower() or rel.content_type.startswith("image/"):
                    att_id = f"image_{len(attachments)}"
                    # Store as relationship reference for now
                    attachments.append(
                        Attachment(
                            id=att_id,
                            path=rel.target_ref,
                            mime_type=rel.content_type,
                        )
                    )
                    logger.debug(f"[parsing:python-docx] Attachment found: {att_id}")
        except Exception:
            logger.debug("[parsing:python-docx] Could not extract image references from relationships")

        logger.info(f"[parsing:python-docx] Complete: {len(blocks)} blocks, {len(attachments)} attachments")
        return IntermediateResult(
            blocks=blocks,
            metadata=metadata,
            attachments=attachments,
        )

    def process(
        self,
        docx_path: Path,
        output_dir: Path,
        config: ProcessorConfig | None = None,
        skip_image_ocr: bool = False,
    ) -> ProcessResult:
        """Parse and process a DOCX file into rich Markdown + placeholder Markdown + resources.

        Calls parse() internally, then uses MarkdownGenerator to produce both
        rich Markdown output and placeholder Markdown with XML tags,
        resource directory, and JSON mapping.

        Args:
            docx_path: Absolute path to the DOCX file.
            output_dir: Directory to write output files to.
            config: Processor configuration options.

        Returns:
            ProcessResult with paths to the generated files.
        """
        if config is None:
            config = ProcessorConfig()

        # Parse to get intermediate result
        intermediate = self.parse(docx_path)

        # Generate both Markdown outputs
        gen = MarkdownGenerator(
            intermediate,
            output_resources=config.output_resources,
        )

        rich_md_path, placeholder_md_path, resources_dir, json_path = gen.generate_both(
            output_dir,
            source_path=str(docx_path),
            parser_name=self.name,
        )

        return ProcessResult(
            rich_markdown_path=rich_md_path if config.output_markdown else None,
            placeholder_markdown_path=placeholder_md_path if config.output_markdown else None,
            resources_dir=resources_dir,
            resources_json_path=json_path,
            metadata=dict(intermediate.metadata),
        )


# Auto-register
ParserRegistry.register("python-docx", PythonDocxParser)
