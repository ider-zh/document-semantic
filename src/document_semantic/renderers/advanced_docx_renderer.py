from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Optional, Dict

import docx
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from document_semantic.models.annotated_content import AnnotatedMinerUContentList
from document_semantic.models.mineru_content import (
    MinerUParagraphContent,
    MinerUTitleContent,
    MinerUInlineContent
)
from document_semantic.templates.schema import SemanticTemplate, DocxStyleConfig

logger = logging.getLogger(__name__)


class AdvancedDocxRenderer:
    """Renders annotated content using detailed YAML-based style and layout definitions."""

    def __init__(self, resources_dir: Optional[Path] = None):
        self._doc = None
        self.resources_dir = resources_dir

    def render(
        self, 
        annotated_content: AnnotatedMinerUContentList, 
        template: SemanticTemplate, 
        output_path: Path
    ) -> Path:
        self._doc = docx.Document()
        
        # 1. Setup Page Layout
        self._setup_page(template)
        
        # 2. Create Styles
        self._setup_styles(template)
        
        # 3. Render Content with Section Management
        current_section_type = "title_page" # title_page -> abstract_page -> body
        
        # Counters for numbering
        fig_count = 0
        tbl_count = 0
        eq_count = 0

        for ann_elem in annotated_content:
            tag = ann_elem.semantic_tag
            elem = ann_elem.element

            # Section switching logic
            new_section_type = self._determine_section_type(tag, current_section_type)
            if new_section_type != current_section_type:
                self._add_section_break(template, new_section_type)
                current_section_type = new_section_type

            # Element specific rendering
            if elem.type == "image":
                # Only count images that have captions for figure numbering
                has_caption = hasattr(elem.content, "image_caption") and elem.content.image_caption
                if has_caption:
                    fig_count += 1
                self._render_image(elem, tag, template, fig_count)
            elif elem.type == "equation_interline":
                eq_count += 1
                self._render_equation(elem, tag, template, eq_count)
            elif elem.type == "table":
                tbl_count += 1
                self._render_table(elem, tag, template, tbl_count)
            else:
                self._render_element(ann_elem, template)

        self._doc.save(str(output_path))
        logger.info(f"[renderer:advanced_docx] Saved document to {output_path}")
        return output_path

    def _render_image(self, elem, tag: str, template: SemanticTemplate, count: int):
        """Render an image element.

        The count parameter is the sequential figure number counter,
        but we only assign figure numbers to images that have captions.
        """
        style_name = tag if tag in template.styles else "Normal"
        img_path = None

        # Try to find actual image file
        if self.resources_dir and hasattr(elem.content, "image_source") and elem.content.image_source:
            # MinerU image path is usually relative to the ZIP root, e.g., 'images/abc.jpg'
            # Our resources_dir should be the parent of 'images'
            rel_path = elem.content.image_source.path
            img_path = self.resources_dir / rel_path

        p = self._doc.add_paragraph(style=style_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if img_path and img_path.exists():
            try:
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(3.2))
            except Exception as e:
                logger.warning(f"[renderer:advanced_docx] Failed to add image: {e}")
                p.add_run(f"[Image: {img_path.name}]")
        else:
            if img_path:
                p.add_run(f"[Image not found: {img_path.name}]")
            else:
                p.add_run("[Image: no source]")

        # Handle Caption - only render if there's a caption
        caption_text = ""
        if hasattr(elem.content, "image_caption") and elem.content.image_caption:
            caption_text = self._extract_inline_list_text(elem.content.image_caption)

        if caption_text:
            # Extract figure number from caption (e.g., "Fig.1: ..." or "Fig. 1: ...")
            import re
            match = re.search(r'^\s*Fig\.?\s*\d+[:\s]*', caption_text, re.IGNORECASE)
            if match:
                # Caption already starts with figure number, use it as-is
                display_caption = caption_text
            else:
                # No figure number in caption, prepend sequential count
                fig_num = str(count)
                display_caption = f"Fig. {fig_num}. {caption_text}"

            cap_p = self._doc.add_paragraph(style="figure_caption" if "figure_caption" in template.styles else style_name)
            cap_p.add_run(display_caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_equation(self, elem, tag: str, template: SemanticTemplate, count: int):
        style_name = tag if tag in template.styles else "Normal"

        # Render equation as OMML (native Word math), not image
        # Full-width centered paragraph for the equation
        p_eq = self._doc.add_paragraph(style=style_name)
        p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

        latex = getattr(elem.content, "math_content", "")
        if latex:
            try:
                from document_semantic.utils.mathml_to_omml import insert_omml_block
                insert_omml_block(p_eq, latex)
            except Exception as e:
                logger.warning(f"[renderer:advanced_docx] OMML conversion failed, falling back to text: {e}")
                p_eq.add_run(latex)
        else:
            p_eq.add_run("[equation]")

        # Equation number in a separate right-aligned paragraph with no spacing
        p_num = self._doc.add_paragraph(style=style_name)
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_num.paragraph_format.space_before = Pt(0)
        p_num.paragraph_format.space_after = Pt(6)
        p_num.add_run(f"({count})")

    def _render_table(self, elem, tag: str, template: SemanticTemplate, count: int):
        style_name = tag if tag in template.styles else "Normal"
        
        # 1. Render Caption first (Table captions are often above)
        caption_text = ""
        if hasattr(elem.content, "table_caption") and elem.content.table_caption:
            caption_text = self._extract_inline_list_text(elem.content.table_caption)
        
        if caption_text:
            cap_p = self._doc.add_paragraph(style="table_caption" if "table_caption" in template.styles else style_name)
            cap_p.add_run(f"Table {count}. {caption_text}")
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. Table content - Parse HTML and create Word table
        html_content = getattr(elem.content, "html", None)
        if html_content:
            try:
                self._render_html_table(html_content, template)
            except Exception as e:
                logger.error(f"[renderer:advanced_docx] Failed to render HTML table: {e}")
                self._doc.add_paragraph(f"[Table {count} Rendering Failed]", style="Normal")
        else:
            self._doc.add_paragraph(f"[Complex Table {count}]", style="Normal")

    def _render_html_table(self, html: str, template: SemanticTemplate):
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "lxml")
        table_tag = soup.find("table")
        if not table_tag:
            return

        rows = table_tag.find_all("tr")
        if not rows:
            return

        # Determine dimensions
        num_rows = len(rows)
        num_cols = 0
        for row in rows:
            cols = row.find_all(["td", "th"])
            num_cols = max(num_cols, len(cols))

        if num_cols == 0:
            return

        # Create table in Word
        docx_table = self._doc.add_table(rows=num_rows, cols=num_cols)
        docx_table.style = "Table Grid" # Base style
        docx_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fill data
        for r_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for c_idx, cell in enumerate(cells):
                if c_idx < num_cols:
                    text = cell.get_text(strip=True)
                    docx_cell = docx_table.cell(r_idx, c_idx)
                    docx_cell.text = text
                    
                    # Apply cell font size if needed
                    for p in docx_cell.paragraphs:
                        p.style = self._doc.styles["Normal"]
                        if p.runs:
                            p.runs[0].font.size = Pt(9) # Smaller font for tables

        # Apply Three-Line Table Style
        self._set_three_line_style(docx_table)

    def _set_three_line_style(self, table):
        """Applies academic three-line style to a docx table."""
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        if num_rows == 0 or num_cols == 0:
            return

        # Remove inner borders first
        for row in table.rows:
            for cell in row.cells:
                self._remove_cell_borders(cell)
        
        # Apply borders
        for c in range(num_cols):
            # Top border of first row
            self._set_cell_border(table.cell(0, c), top={"sz": 12, "val": "single", "color": "000000"})
            # Bottom border of first row (header separator)
            self._set_cell_border(table.cell(0, c), bottom={"sz": 6, "val": "single", "color": "000000"})
            # Bottom border of last row
            self._set_cell_border(table.cell(num_rows - 1, c), bottom={"sz": 12, "val": "single", "color": "000000"})

    def _remove_cell_borders(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            # Insert before w:shd if present, else append
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                shd.addprevious(tcBorders)
            else:
                tcPr.append(tcBorders)
        
        for edge in ('top', 'left', 'bottom', 'right', 'tl2br', 'tr2bl'):
            edge_tag = qn(f'w:{edge}')
            element = tcBorders.find(edge_tag)
            if element is None:
                element = parse_xml(f'<w:{edge} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tcBorders.append(element)
            element.set(qn('w:val'), 'nil')

    def _set_cell_border(self, cell, **kwargs):
        """
        Set cell borders.
        Usage: _set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"}, ...)
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            # Insert before w:shd if present, else append
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                shd.addprevious(tcBorders)
            else:
                tcPr.append(tcBorders)

        for edge, props in kwargs.items():
            if edge in ('top', 'left', 'bottom', 'right', 'tl2br', 'tr2bl'):
                edge_tag = qn(f'w:{edge}')
                element = tcBorders.find(edge_tag)
                if element is None:
                    element = parse_xml(f'<w:{edge} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    tcBorders.append(element)
                
                for key, val in props.items():
                    element.set(qn(f'w:{key}'), str(val))

    def _extract_inline_list_text(self, inlines: list) -> str:
        return "".join([i.content for i in inlines if hasattr(i, "content")]).strip()

    def _setup_page(self, template: SemanticTemplate):
        section = self._doc.sections[0]
        page = template.page
        
        section.page_width = self._parse_unit(page.width)
        section.page_height = self._parse_unit(page.height)
        section.orientation = WD_ORIENT.PORTRAIT if page.orientation == "portrait" else WD_ORIENT.LANDSCAPE
        
        section.top_margin = self._parse_unit(page.margins.top)
        section.bottom_margin = self._parse_unit(page.margins.bottom)
        section.left_margin = self._parse_unit(page.margins.left)
        section.right_margin = self._parse_unit(page.margins.right)
        
        # Initial columns (title page)
        self._set_columns(section, page.columns.title_page, page.columns.column_spacing)

    def _setup_styles(self, template: SemanticTemplate):
        styles = self._doc.styles
        for name, config in template.styles.items():
            # Create a user-defined style if it doesn't exist
            # Note: We use the internal name which might clash with Word defaults
            # To be safe, we can prefix them, but let's try direct for now.
            try:
                style = styles.add_style(name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            except ValueError:
                style = styles[name] # Already exists

            self._apply_style_config(style, config)

    def _apply_style_config(self, style, config: DocxStyleConfig):
        font = style.font
        if config.font:
            font.name = config.font
        if config.size:
            font.size = Pt(config.size)
        font.bold = config.bold
        font.italic = config.italic
        
        par_fmt = style.paragraph_format
        par_fmt.alignment = self._parse_alignment(config.align)
        par_fmt.space_before = Pt(config.space_before)
        par_fmt.space_after = Pt(config.space_after)
        par_fmt.line_spacing = config.line_spacing
        # Only set indentation if explicitly configured (non-zero)
        # Setting Pt(0) can cause unwanted hanging indents in python-docx
        if config.indent_first > 0:
            par_fmt.first_line_indent = Pt(config.indent_first)
        if config.indent_left > 0:
            par_fmt.left_indent = Pt(config.indent_left)

    def _determine_section_type(self, tag: str, current: str) -> str:
        if tag in ("paper_title", "author_info"):
            return "title_page"
        if tag in ("abstract_head", "abstract_text", "index_terms"):
            return "abstract_page"
        return "body"

    def _add_section_break(self, template: SemanticTemplate, section_type: str):
        new_section = self._doc.add_section(WD_SECTION.CONTINUOUS)
        cols = template.page.columns
        count = getattr(cols, section_type, 1)
        self._set_columns(new_section, count, cols.column_spacing)

    def _set_columns(self, section, count: int, spacing: str):
        sectPr = section._sectPr
        cols = sectPr.xpath("./w:cols")
        if not cols:
            cols = parse_xml(f'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />')
            # Insert before elements that must come after w:cols (vAlign, docGrid, etc.)
            vAlign = sectPr.find(qn("w:vAlign"))
            docGrid = sectPr.find(qn("w:docGrid"))
            if vAlign is not None:
                vAlign.addprevious(cols)
            elif docGrid is not None:
                docGrid.addprevious(cols)
            else:
                sectPr.append(cols)
        else:
            cols = cols[0]
            
        cols.set(qn("w:num"), str(count))
        cols.set(qn("w:space"), str(int(self._parse_unit(spacing).twips)))

    def _render_element(self, ann_elem, template: SemanticTemplate):
        tag = ann_elem.semantic_tag
        elem = ann_elem.element
        style_name = tag if tag in template.styles else "Normal"

        if elem.type in ("title", "paragraph"):
            p = self._doc.add_paragraph(style=style_name)
            inlines = []
            if elem.type == "title":
                inlines = elem.content.title_content or []
            else:
                inlines = elem.content.paragraph_content or []

            self._render_inlines(p, inlines)

            # Apply text transform
            config = template.styles.get(tag)
            if config and config.text_transform == "uppercase":
                for run in p.runs:
                    run.text = run.text.upper()
        
        elif elem.type == "list":
            # Render list items as paragraphs for reference items
            for list_item in elem.content.list_items:
                p = self._doc.add_paragraph(style=style_name)
                self._render_inlines(p, list_item.item_content)

                # Apply text transform
                config = template.styles.get(tag)
                if config and config.text_transform == "uppercase":
                    for run in p.runs:
                        run.text = run.text.upper()

        elif elem.type == "equation_interline":
            p = self._doc.add_paragraph(elem.content.math_content, style=style_name)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_inlines(self, paragraph, inlines: list):
        """Render a list of inline elements into a paragraph."""
        for inline in inlines:
            if inline.type == "equation_inline":
                # Render as native OMML inline equation from LaTeX source
                try:
                    from document_semantic.utils.mathml_to_omml import insert_omml_inline
                    insert_omml_inline(paragraph, inline.content)
                except Exception as e:
                    logger.warning(f"[renderer:advanced_docx] Inline OMML failed, falling back to text: {e}")
                    run = paragraph.add_run(f"${inline.content}$")
                    run.font.name = "Cambria Math"
            elif inline.type == "text":
                text = inline.content
                # Basic cleanup of markdown-style links/images if they leaked in
                text = re.sub(r"!\[.*?\]\((.*?)\)", r"\1", text)
                text = re.sub(r"\[.*?\]\((.*?)\)", r"\1", text)
                paragraph.add_run(text)

    def _parse_unit(self, value: str):
        if value.endswith("mm"):
            return Mm(float(value[:-2]))
        if value.endswith("pt"):
            return Pt(float(value[:-2]))
        if value.endswith("in"):
            return Inches(float(value[:-2]))
        return Pt(float(value))

    def _parse_alignment(self, align: str):
        return {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justified": WD_ALIGN_PARAGRAPH.JUSTIFY
        }.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)
