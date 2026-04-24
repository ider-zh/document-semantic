from __future__ import annotations

import logging
from pathlib import Path
from typing import List, Optional

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from document_semantic.models.annotated_content import AnnotatedMinerUContentList, AnnotatedMinerUElement
from document_semantic.models.mineru_content import (
    MinerUParagraphContent,
    MinerUTitleContent,
    MinerUImageContent,
    MinerUInlineContent
)
from document_semantic.templates.schema import SemanticTemplate

logger = logging.getLogger(__name__)


class DocxRenderer:
    """Renders annotated MinerU content to a DOCX file using template styles."""

    def __init__(self, template_path: Optional[Path] = None):
        """
        Args:
            template_path: Path to a .docx file with pre-defined styles (JCST/IEEE).
        """
        self.template_path = template_path

    def render(
        self, 
        annotated_content: AnnotatedMinerUContentList, 
        template: SemanticTemplate, 
        output_path: Path
    ) -> Path:
        """Renders the content to a DOCX file."""
        doc = docx.Document(str(self.template_path) if self.template_path else None)
        
        style_map = template.docx_style_map
        
        for ann_elem in annotated_content:
            tag = ann_elem.semantic_tag
            elem = ann_elem.element
            style_name = style_map.get(tag, "Normal")
            
            try:
                if elem.type == "title":
                    # For titles, MinerU has level. We can use style_name or default Heading X
                    p = doc.add_paragraph(style=style_name)
                    self._add_inline_content(p, elem.content.title_content or [])
                    
                elif elem.type == "paragraph":
                    p = doc.add_paragraph(style=style_name)
                    self._add_inline_content(p, elem.content.paragraph_content or [])
                    
                elif elem.type == "image":
                    # Placeholder for images. In real case, we need the local image path.
                    doc.add_paragraph(f"[Image Placeholder: {tag}]", style=style_name)
                    
                elif elem.type == "equation_interline":
                    # Add as a centered paragraph or specialized style
                    p = doc.add_paragraph(elem.content.math_content, style=style_name)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                else:
                    # Fallback
                    text = self._extract_text_fallback(elem)
                    if text:
                        doc.add_paragraph(text, style=style_name)
            except Exception as e:
                logger.error(f"[renderer:docx] Failed to render element {elem.type} with style {style_name}: {e}")
                # Fallback to plain Normal
                doc.add_paragraph(self._extract_text_fallback(elem))

        doc.save(str(output_path))
        logger.info(f"[renderer:docx] Saved document to {output_path}")
        return output_path

    def _add_inline_content(self, paragraph, inlines: List[MinerUInlineContent]):
        for inline in inlines:
            run = paragraph.add_run(inline.content)
            if inline.type == "equation_inline":
                # For inline equations in DOCX, it is complex. 
                # For now, just keep as text or use a specific font.
                run.italic = True 

    def _extract_text_fallback(self, elem) -> str:
        parts = []
        def _collect(obj):
            if isinstance(obj, str): parts.append(obj)
            elif isinstance(obj, dict):
                for k, v in obj.items():
                    if k == "content": _collect(v)
                    elif isinstance(v, (str, dict, list)): _collect(v)
            elif isinstance(obj, list):
                for x in obj: _collect(x)
        _collect(elem.content)
        return "".join(parts).strip()
