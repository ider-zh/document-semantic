"""Inspect the full_test.docx structure."""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import json

doc = Document('asset/full_test.docx')

info = {
    'paragraph_count': len(doc.paragraphs),
    'paragraphs': [],
    'table_count': len(doc.tables),
    'section_count': len(doc.sections),
}

for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text_preview = p.text[:100] if p.text else '(empty)'
    
    # Check for XML elements (column breaks, etc.)
    has_break = False
    for run in p.runs:
        for child in run._element:
            if child.tag.endswith('br'):
                br_type = child.get(qn('w:type'), 'none')
                break_attr = child.get(qn('w:colBefore'), 'none')
                has_break = True
    
    # Check for inline images
    has_image = any(
        child.tag.endswith('drawing') or child.tag.endswith('imagedata')
        for run in p.runs
        for child in run._element
    )
    
    info['paragraphs'].append({
        'idx': i,
        'style': style,
        'text': text_preview,
        'has_break': has_break,
        'has_image': has_image,
        'run_count': len(p.runs),
    })

for s in doc.sections:
    info['sections'].append({
        'gutter': str(s.gutter),
        'page_width': str(s.page_width),
        'page_height': str(s.page_height),
        'orientation': 'landscape' if s.page_width > s.page_height else 'portrait',
    })

# Check for images at document level
rels = doc.part.rels
info['doc_rels'] = {}
for k, v in rels.items():
    info['doc_rels'][k] = {
        'type': v.reltype,
        'target': v.target_ref,
    }

with open('_inspect_output.json', 'w', encoding='utf-8') as f:
    json.dump(info, f, indent=2, ensure_ascii=False)

# Also print summary for quick view
for p in info['paragraphs']:
    marker = ''
    if p['has_break']:
        marker += ' [BREAK]'
    if p['has_image']:
        marker += ' [IMAGE]'
    print(f"  [{p['idx']:2d}] {p['style']:30s} {p['text']!r}{marker}")

print(f"\nSections: {len(info.get('sections', []))}")
print(f"Doc-level image relationships: {len(info['doc_rels'])}")
